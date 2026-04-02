import openai
import time
import re
import json
import logging
import datetime
from django.utils.html import strip_tags
from django.utils import timezone
from django.db.models import Sum
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from io import BytesIO
from prompts_app.services import get_active_prompt_for_job
from prompts_app.models import Prompt
from .prompt_strings import (
    DEFAULT_SYSTEM_PROMPT,
    BULLETS_SYSTEM_PROMPT,
    BULLETS_PROMPT_BASE,
    BULLETS_PROMPT_NO_BASE,
    BULLETS_REPAIR_PROMPT,
    BULLETS_MISSING_TERMS_PROMPT,
    BUILD_PROMPT_SUMMARY_INSTRUCTION,
    BUILD_PROMPT_SKILLS_WITH_PROFILE,
    BUILD_PROMPT_SKILLS_GENERATE,
    BUILD_PROMPT_BASE_RESUME_LABEL,
    BUILD_PROMPT_BASE_RESUME_MISSING,
    BUILD_PROMPT_BASE_RESUME_GUIDE_1,
    BUILD_PROMPT_BASE_RESUME_GUIDE_2,
    BUILD_PROMPT_BASE_RESUME_GUIDE_3,
    BUILD_PROMPT_EXPERIENCE_LABEL,
    BUILD_PROMPT_NO_EXPERIENCE,
    BUILD_PROMPT_EDUCATION_LABEL,
    BUILD_PROMPT_NO_EDUCATION,
    BUILD_PROMPT_JD_HEADER,
    BUILD_PROMPT_ROLE_NO_DESC,
    BUILD_PROMPT_TEMPLATE_SUFFIX,
    BUILD_PROMPT_BASE_SECTION_WITH,
    BUILD_PROMPT_BASE_SECTION_WITHOUT,
    BUILD_PROMPT_REQUIRED_SECTIONS,
)
from django.utils.html import strip_tags
from core.models import LLMConfig, LLMUsageLog, PlatformConfig
from core.security import decrypt_value
from core.llm_services import calculate_cost

logger = logging.getLogger("apps.resumes")

FILLER_PHRASES = []
EXPANSION_PHRASES = []
NOISE_PHRASES = []

ACTION_VERBS = {
    "improved","reduced","increased","decreased","optimized","streamlined","built","designed","implemented",
    "delivered","automated","migrated","refactored","enhanced","led","managed","supported","resolved","developed",
    "configured","deployed","monitored","troubleshot","maintained","documented","collaborated","coordinated",
}

def _fill_prompt(template, **kwargs):
    """Replace {key} placeholders without interpreting other braces."""
    text = template
    for k, v in kwargs.items():
        text = text.replace("{" + k + "}", v)
    return text



STOPWORDS = {
    "the","and","for","with","that","this","from","your","you","our","are","was","were","will","shall","can",
    "able","ability","have","has","had","not","but","use","using","used","into","over","under","across","per",
    "to","of","in","on","at","by","as","or","an","a","is","it","we","they","their","them","he","she","his","her",
    "be","been","being","if","then","than","also","such","other","more","most","less","least","any","all","each",
    "including","include","includes","within","without","via","etc","etc.",
    # JD noise words — locations, generic descriptors, filler
    "amazon","knowledge","familiarity","general","concepts","understanding","willingness",
    "demonstrated","mastery","preferred","required","excellent","passion","must","focus",
    "experience","basic","least","jersey","city","holmdel","states","united","remote",
    "hybrid","location","responsibilities","qualifications","description","job","role",
    "candidate","ideal","looking","seeking","opportunity","team","company","organization",
    "environment","position","based","work","working","year","years","new","would","like",
    "should","well","good","great","strong","highly","about","need","needs","ensure",
    "minimum","maximum","salary","salaries","commission","compensation","benefits","posting",
    "schedule","full","time","apply","applynow","date","identification","id",
}


def extract_keywords(text, max_keywords=200):
    if not text:
        return []
    tokens = re.findall(r"[A-Za-z0-9][A-Za-z0-9+./#_-]{1,}", text.lower())
    keywords = []
    for t in tokens:
        if t in {
            "disclaimer","privacy","notice","sti","lti","apply","applynow","hybrid","remote",
            "job","identification","id","category","information","technology","posting","date",
            "schedule","locations","location","division","employer","legal","annualized","base",
            "pay","salary","commission","work","arrangement","full","time","yes","no","now",
        }:
            continue
        if len(t) < 3:
            continue
        if t.isdigit():
            continue
        if t in STOPWORDS:
            continue
        keywords.append(t)
    # preserve order, unique
    seen = set()
    uniq = []
    for k in keywords:
        if k in seen:
            continue
        seen.add(k)
        uniq.append(k)
        if len(uniq) >= max_keywords:
            break
    return uniq


def _phrase_tokens(text):
    tokens = re.findall(r"[A-Za-z0-9][A-Za-z0-9+./#_-]{1,}", (text or "").lower())
    return [t for t in tokens if t and t not in STOPWORDS]


def _build_jd_phrase_set(jd_text, n=4):
    tokens = _phrase_tokens(jd_text)
    phrases = set()
    for i in range(len(tokens) - n + 1):
        phrases.add(" ".join(tokens[i:i+n]))
    return phrases


def _contains_jd_long_phrase(line, jd_phrase_set, n=4):
    tokens = _phrase_tokens(line)
    for i in range(len(tokens) - n + 1):
        if " ".join(tokens[i:i+n]) in jd_phrase_set:
            return True
    return False


def _humanize_bullet(line):
    if not line:
        return line
    text = line.strip()
    # soften overly templated phrasing
    text = re.sub(r"\busing\b", "leveraging", text, count=1)
    text = re.sub(r"\bto improve\b", "to strengthen", text, count=1)
    return text


def _apply_jd_alignment_rules(bullets, jd_text, max_keyword_reuse=2):
    jd_keywords = set(extract_keywords(jd_text or "", max_keywords=40))
    jd_phrase_set = _build_jd_phrase_set(jd_text or "", n=4)
    keyword_use = {k: 0 for k in jd_keywords}
    out = []

    # First pass: humanize and count keyword presence
    for b in bullets:
        text = b
        if _contains_jd_long_phrase(text, jd_phrase_set, n=4):
            text = _humanize_bullet(text)
        out.append(text)

    # Keep keywords as-is; do not delete to avoid broken sentences.
    for i, text in enumerate(out):
        out[i] = re.sub(r"\s{2,}", " ", out[i]).strip(" ,.-")

    # No keyword injection or topic appending; rely on LLM instructions.
    return out


def score_ats(jd_text, resume_text):
    if not jd_text or not resume_text:
        return 0
    keywords = extract_keywords(jd_text)
    if not keywords:
        return 0
    content = resume_text.lower()
    matched = [k for k in keywords if k in content]
    score = int((len(matched) / len(keywords)) * 100)
    if score > 100:
        score = 100
    return score


def validate_resume(content):
    errors = []
    warnings = []
    if not content:
        errors.append("Resume content is empty.")
        return errors, warnings

    lines = [l.rstrip() for l in content.splitlines()]
    text = "\n".join(lines)

    required_headings = [
        "PROFESSIONAL SUMMARY",
        "SKILLS",
        "PROFESSIONAL EXPERIENCE",
        "EDUCATION",
    ]

    # Exact heading checks
    for h in required_headings:
        if h not in text:
            errors.append(f"Missing required section heading: {h}.")

    # Enforce section order
    heading_positions = {h: text.find(h) for h in required_headings if h in text}
    if len(heading_positions) == len(required_headings):
        if not (heading_positions["PROFESSIONAL SUMMARY"] <
                heading_positions["SKILLS"] <
                heading_positions["PROFESSIONAL EXPERIENCE"] <
                heading_positions["EDUCATION"]):
            errors.append("Section order must be: PROFESSIONAL SUMMARY, SKILLS, PROFESSIONAL EXPERIENCE, EDUCATION.")

    # Bullet checks in Professional Experience
    if "PROFESSIONAL EXPERIENCE" in text:
        exp_block = text.split("PROFESSIONAL EXPERIENCE", 1)[1]
        for h in ["EDUCATION", "CERTIFICATIONS"]:
            if h in exp_block:
                exp_block = exp_block.split(h, 1)[0]
        lines_exp = [line.rstrip() for line in exp_block.splitlines()]
        bullets = [line for line in lines_exp if line.strip().startswith("- ")]
        if len(bullets) < 6:
            errors.append("PROFESSIONAL EXPERIENCE must include at least 6 bullet points.")

    # Word count per bullet (22–25 words)
    for i, b in enumerate(bullets, start=1):
        words = [w for w in re.findall(r"[A-Za-z0-9']+", b) if w]
        if len(words) < 22:
            errors.append(f"Bullet {i} in PROFESSIONAL EXPERIENCE has fewer than 22 words.")
            break
        if len(words) > 25:
            errors.append(f"Bullet {i} in PROFESSIONAL EXPERIENCE has more than 25 words.")
            break

        # Detect role headers (single-line or two-line with dates)
        header_pattern = re.compile(
            r"^.+\s+\|\s+.+\s+\|\s+.+\d{4}\s*[–-]\s*(Present|\d{4})\s*$"
        )
        roles = []
        i = 0
        while i < len(lines_exp):
            line = lines_exp[i].strip()
            if not line or line.startswith("- "):
                i += 1
                continue
            if header_pattern.match(line):
                roles.append({"start": i})
                i += 1
                continue
            if i + 1 < len(lines_exp):
                next_line = lines_exp[i + 1].strip()
                if next_line and not next_line.startswith("- ") and re.search(r"\d{4}", next_line):
                    roles.append({"start": i})
                    i += 2
                    continue
            i += 1

        if not roles:
            errors.append("Role headers must follow format: Title | Company | Start Date – End Date.")
        else:
            for idx, role in enumerate(roles):
                start = role["start"]
                end = roles[idx + 1]["start"] if idx + 1 < len(roles) else len(lines_exp)
                role_lines = lines_exp[start:end]
                role_bullets = [l for l in role_lines if l.strip().startswith("- ")]
                if idx == 0:
                    if not (7 <= len(role_bullets) <= 10):
                        errors.append(f"Most recent role must have 7–10 bullets (found {len(role_bullets)}).")
                        break
                else:
                    if len(role_bullets) != 6:
                        errors.append(f"Role {idx + 1} must have exactly 6 bullets (found {len(role_bullets)}).")
                        break
    else:
        errors.append("PROFESSIONAL EXPERIENCE section missing or not detected.")

    # Skills format check (key:value lines, no bullets)
    if "SKILLS" in text:
        skills_block = text.split("SKILLS", 1)[1]
        for h in ["PROFESSIONAL EXPERIENCE", "EDUCATION", "CERTIFICATIONS"]:
            if h in skills_block:
                skills_block = skills_block.split(h, 1)[0]
        skills_lines = [l.strip() for l in skills_block.splitlines() if l.strip()]
        if not skills_lines:
            errors.append("SKILLS section is empty.")
        else:
            for line in skills_lines:
                if line.startswith("-"):
                    errors.append("SKILLS must use key:value lines (no bullets).")
                    break
                if ":" not in line:
                    errors.append("SKILLS lines must follow key:value format.")
                    break

    # Summary length check (70–80 words)
    if "PROFESSIONAL SUMMARY" in text:
        summary_block = text.split("PROFESSIONAL SUMMARY", 1)[1]
        for h in ["SKILLS", "PROFESSIONAL EXPERIENCE", "EDUCATION", "CERTIFICATIONS"]:
            if h in summary_block:
                summary_block = summary_block.split(h, 1)[0]
        summary_words = re.findall(r"[A-Za-z0-9']+", summary_block)
        if len(summary_words) < 70 or len(summary_words) > 80:
            errors.append("PROFESSIONAL SUMMARY must be 70–80 words.")
    else:
        errors.append("PROFESSIONAL SUMMARY section missing or not detected.")

    return errors, warnings


def _find_section_bounds(text, heading, headings):
    start = text.find(heading)
    if start == -1:
        return None
    # find next heading after start
    after = text[start + len(heading):]
    next_positions = []
    for h in headings:
        if h == heading:
            continue
        idx = after.find(h)
        if idx != -1:
            next_positions.append(idx)
    end = start + len(heading) + (min(next_positions) if next_positions else len(after))
    return start, end


def extract_section(content, heading, headings):
    bounds = _find_section_bounds(content, heading, headings)
    if not bounds:
        return ""
    start, end = bounds
    return content[start:end].strip()


def replace_section(content, heading, headings, new_section):
    bounds = _find_section_bounds(content, heading, headings)
    if not bounds:
        return content
    start, end = bounds
    return (content[:start] + new_section.strip() + "\n\n" + content[end:]).strip()


def _format_month_year(dt):
    if not dt:
        return ""
    try:
        return dt.strftime('%b %Y')
    except Exception:
        return str(dt)


def _build_header_block(job, consultant):
    name = consultant.user.get_full_name() or consultant.user.username
    # Extract city, state from JD location (e.g. "Holmdel, NJ, United States" → "Holmdel, NJ")
    location = _extract_city_state(job.location) or "United States"
    phone = consultant.phone or ""
    email = consultant.user.email or ""
    contact_parts = []
    if location:
        contact_parts.append(location)
    if email:
        contact_parts.append(email)
    if phone:
        contact_parts.append(phone)
    contact_line = " | ".join(contact_parts)
    return f"{name}\n{contact_line}".strip()


def _extract_city_state(location_text):
    """Extract 'City, ST' from a location string like 'Holmdel, NJ, United States'."""
    if not location_text:
        return ""
    # US state abbreviations
    us_states = {
        'AL','AK','AZ','AR','CA','CO','CT','DE','FL','GA','HI','ID','IL','IN','IA',
        'KS','KY','LA','ME','MD','MA','MI','MN','MS','MO','MT','NE','NV','NH','NJ',
        'NM','NY','NC','ND','OH','OK','OR','PA','RI','SC','SD','TN','TX','UT','VT',
        'VA','WA','WV','WI','WY','DC',
    }
    parts = [p.strip() for p in location_text.split(',')]
    # Find the state abbreviation
    city = parts[0] if parts else ""
    state = ""
    for p in parts[1:]:
        p_clean = p.strip().upper()
        if p_clean in us_states:
            state = p_clean
            break
    if city and state:
        return f"{city}, {state}"
    # If no US state found, return first two parts
    if len(parts) >= 2:
        return f"{parts[0]}, {parts[1]}"
    return location_text.strip()


def _normalize_match_text(text):
    return re.sub(r"[^a-z0-9]+", "", (text or "").lower())

def _clean_bullet_line(line):
    if not line:
        return ""
    line = line.strip()
    line = re.sub(r"^[-•*]\s+", "", line)
    line = re.sub(r"^\d+[\.\)]\s+", "", line)
    line = re.sub(r"^\*\*", "", line)
    line = re.sub(r"\*\*$", "", line)
    return line.strip()


def _is_skill_bullet(line):
    if not line:
        return False
    normalized = _clean_bullet_line(line).lower()
    if ":" not in normalized:
        return False
    label, items = [p.strip() for p in normalized.split(":", 1)]
    if not label or not items:
        return False
    # Heuristic: skill lines usually have short labels and comma-separated items.
    if len(label.split()) > 4:
        return False
    item_parts = [p.strip() for p in items.split(",") if p.strip()]
    return len(item_parts) >= 2


def _bullet_word_count(line):
    return len([w for w in re.findall(r"[A-Za-z0-9']+", line) if w])


def _expand_bullet_to_min_words_strict(line, job, min_words, method_keywords):
    if not line:
        return line
    # Do NOT inject extra keywords; rely on LLM to meet length.
    return line.strip().rstrip(".") + "."


def _cap_bullet_words(line, max_words=25):
    if not line:
        return line
    words = re.findall(r"[A-Za-z0-9']+", line)
    if len(words) <= max_words:
        return line.strip().rstrip(".") + "."
    tokens = line.split()
    trimmed = " ".join(tokens[:max_words])
    if not trimmed.endswith("."):
        trimmed = trimmed.rstrip(" ,;-") + "."
    return trimmed

def _strip_vague_adverbs(line):
    if not line:
        return line
    adverbs = {"significantly", "substantially", "notably", "dramatically", "remarkably"}
    tokens = line.split()
    filtered = [t for t in tokens if t.lower().strip(",.;") not in adverbs]
    out = " ".join(filtered)
    return out.strip()

def _normalize_bullet_for_dedupe(line):
    if not line:
        return ""
    text = _clean_bullet_line(line).lower()
    text = re.sub(r"[^a-z0-9\s]+", "", text)
    words = [w for w in text.split() if w and w not in STOPWORDS]
    return " ".join(words)


def _dedupe_bullets(lines):
    seen = set()
    out = []
    for b in lines:
        norm = _normalize_bullet_for_dedupe(b)
        if not norm:
            continue
        if norm in seen:
            continue
        seen.add(norm)
        out.append(b)
    return out


def _jd_requires_metrics(job):
    text = (job.description or "").lower()
    signals = [
        "kpi", "sla", "slo", "uptime", "availability", "mttr", "latency",
        "throughput", "performance", "optimiz", "cost", "savings", "reduction",
        "efficiency", "benchmark", "baseline", "roi",
    ]
    return any(s in text for s in signals)


def _collect_method_keywords(job, consultant):
    keywords = set()
    for s in (consultant.skills or []):
        s = s.strip().lower()
        if len(s) >= 3:
            keywords.add(s)
    for k in extract_keywords(job.description or "", max_keywords=30):
        if len(k) >= 4:
            keywords.add(k.lower())
    return keywords


def _clean_jd_text(jd_text):
    if not jd_text:
        return ""
    lines = jd_text.splitlines()
    drop_patterns = [
        r"job identification", r"job category", r"posting date", r"job schedule", r"locations",
        r"sti", r"lti", r"commission", r"work arrangement", r"minimum salary", r"maximum salary",
        r"division", r"legal employer", r"disclaimer", r"apply now",
        r"verisk", r"great place to work", r"equal opportunity", r"employee privacy notice",
    ]
    cleaned = []
    for line in lines:
        low = line.strip().lower()
        if not low:
            cleaned.append(line)
            continue
        if any(re.search(p, low) for p in drop_patterns):
            continue
        if re.match(r"^\d{2}/\d{2}/\d{4}", low):
            continue
        if re.match(r"^\d+$", low):
            continue
        cleaned.append(line)
    return "\n".join(cleaned).strip()


def _bullet_has_action(line):
    if not line:
        return False
    first = _clean_bullet_line(line).split(" ", 1)[0].lower()
    return first in ACTION_VERBS


def _bullet_has_method(line, method_keywords):
    if not line or not method_keywords:
        return False
    low = line.lower()
    return any(k in low for k in method_keywords)


def _bullet_has_metric(line):
    if not line:
        return False
    if re.search(r"\$\s?\d", line):
        return True
    if re.search(r"\d+(\.\d+)?\s?%", line):
        return True
    if re.search(r"\b\d+(\.\d+)?\s?(ms|s|sec|secs|second|seconds|min|mins|minute|minutes|hour|hours|day|days|x)\b", line, re.I):
        return True
    if re.search(r"\b\d{3,}\b", line):
        return True
    return False


def _fix_broken_bullet_text(line):
    """Fix common LLM output defects (orphaned %, ~, broken grammar)."""
    if not line:
        return line
    out = line
    # Orphaned "by%" or "by %" (no number before %)
    out = re.sub(r"\bby\s*%", "significantly", out, flags=re.I)
    # Orphaned bare "%" not preceded by a number
    out = re.sub(r"(?<!\d)\s*%", "", out)
    # "ensuring ~ data availability" → "ensuring high data availability" (BEFORE generic ~ removal)
    out = re.sub(r"~\s*(data|system|service|network)", r"high \1", out, flags=re.I)
    # Orphaned "~" not followed by a number (catch-all, AFTER specific ~ patterns above)
    out = re.sub(r"~\s*(?!\d)", "", out)
    # "a increase" → "a significant increase" (missing adjective)
    out = re.sub(r"\ba\s+(increase|decrease|reduction|improvement)\b", r"a significant \1", out, flags=re.I)
    # Trailing "using solutions" (generic LLM filler)
    out = re.sub(r"\s+using\s+solutions\.?\s*$", ".", out, flags=re.I)
    # Double spaces
    out = re.sub(r"\s{2,}", " ", out).strip()
    # Ensure ends with period
    if out and not out.endswith("."):
        out = out.rstrip(" ,;-") + "."
    return out


def _strip_metrics(line):
    """Remove metric numbers from a bullet while preserving readable grammar."""
    if not line:
        return line
    out = re.sub(r"\$\s?\d+([,.\d]+)?", "", line)
    # "by 25%" → "significantly" (remove orphaned 'by' with number)
    out = re.sub(r"\bby\s+\d+([,.\d]+)?\s?%", "significantly", out, flags=re.I)
    # "achieving a 40% reduction" → "achieving a significant reduction"
    out = re.sub(r"\ba\s+\d+([,.\d]+)?\s?%", "a significant", out, flags=re.I)
    # "~\s?number" → remove
    out = re.sub(r"~\s?\d+([,.\d]+)?\s?%?", "", out)
    # Remaining bare percentages
    out = re.sub(r"\d+([,.\d]+)?\s?%", "", out)
    # Time-based metrics
    out = re.sub(r"\b\d+([,.\d]+)?\s?(ms|s|sec|secs|second|seconds|min|mins|minute|minutes|hour|hours|day|days|x)\b", "", out, flags=re.I)
    # Large numbers
    out = re.sub(r"\b\d{3,}\b", "", out)
    # Clean up broken grammar artifacts
    out = re.sub(r"\bby\s*,", ",", out)  # orphaned "by,"
    out = re.sub(r"\bby\s*$", "", out)  # trailing "by"
    out = re.sub(r"\bby\s+\.", ".", out)  # "by ."
    out = re.sub(r"\s*,\s*,", ",", out)  # double commas
    out = re.sub(r"\s{2,}", " ", out).strip(" ,.-")
    return out


def _round_metrics(line):
    if not line:
        return line
    def repl_percent(m):
        num = float(m.group(1))
        return f"~{int(round(num))}%"
    out = re.sub(r"\b(\d+\.\d+)\s?%", repl_percent, line)
    def repl_decimal(m):
        num = float(m.group(1))
        return f"~{int(round(num))}"
    out = re.sub(r"\b(\d+\.\d+)\b", repl_decimal, out)
    def repl_large(m):
        num = int(m.group(0))
        if num >= 10000:
            return f"over {num//1000}k"
        return str(num)
    out = re.sub(r"\b\d{4,}\b", repl_large, out)
    return out


def _cap_metrics_for_role(job, role_index):
    if _jd_requires_metrics(job):
        return 2 if role_index == 0 else 1
    return 1 if role_index == 0 else 0


def _apply_metric_rules(bullets, job, method_keywords, max_metrics):
    cleaned = []
    metrics_used = 0
    for b in bullets:
        text = b
        # remove vague metric words if no actual metric
        if not _bullet_has_metric(text):
            text = re.sub(r"\b(significant|significantly|notable|substantial)\b", "", text, flags=re.I).strip(" ,.-")
        has_metric = _bullet_has_metric(text)
        if has_metric:
            text = _round_metrics(text)
            has_action = _bullet_has_action(text)
            has_method = _bullet_has_method(text, method_keywords)
            if not (has_action and has_method):
                text = _strip_metrics(text)
                has_metric = _bullet_has_metric(text)
        if has_metric:
            if metrics_used >= max_metrics:
                text = _strip_metrics(text)
                has_metric = _bullet_has_metric(text)
            else:
                metrics_used += 1
        cleaned.append(text)
    return cleaned


def _expand_bullet_to_min_words(line, job, min_words=22):
    """Expand short bullets to meet minimum word count by adding
    relevant JD context. Uses proper grammar, not garbage filler."""
    if not line:
        return line
    words = _bullet_word_count(line)
    if words >= min_words:
        # Trim if over max (32 words)
        if words > 32:
            word_list = line.split()
            line = " ".join(word_list[:32])
            if not line.endswith("."):
                line = line.rstrip(" ,;-") + "."
        return line
    return line


def _total_experience_years_display(consultant):
    experiences = list(consultant.experience.all())
    if not experiences:
        return None
    starts = [e.start_date for e in experiences if e.start_date]
    if not starts:
        return None
    earliest = min(starts)
    latest = None
    current = any(e.is_current for e in experiences)
    for e in experiences:
        if e.is_current:
            latest = datetime.date.today()
            break
    if not latest:
        ends = [e.end_date for e in experiences if e.end_date]
        latest = max(ends) if ends else datetime.date.today()
    months = (latest.year - earliest.year) * 12 + (latest.month - earliest.month)
    years = max(0, months // 12)
    if current:
        return f"{years}+"
    return str(years)


def _required_terms_from_jd(jd_text):
    return extract_required_terms_from_jd(jd_text)


def _normalize_core_skills_format(text):
    if "SKILLS" not in text:
        return text
    headings = ["PROFESSIONAL SUMMARY", "SKILLS", "PROFESSIONAL EXPERIENCE", "CERTIFICATIONS", "EDUCATION"]
    section = extract_section(text, "SKILLS", headings)
    if not section:
        return text
    lines = section.splitlines()
    normalized = []
    for line in lines:
        m = re.match(r"^\s*[-•*]?\s*([^:]+):\s*(.*)$", line.strip())
        if m:
            label = m.group(1).strip()
            items = m.group(2).strip()
            normalized.append(f"{label}: {items}".strip())
        else:
            normalized.append(line.strip())
    section = "\n".join([l for l in normalized if l])
    return replace_section(text, "SKILLS", headings, section)


def _extract_bullets_for_role(content, title, company):
    if not content or not title or not company:
        return []
    lines = [l.rstrip() for l in content.splitlines()]
    title_l = title.lower()
    company_l = company.lower()
    title_n = _normalize_match_text(title)
    company_n = _normalize_match_text(company)
    start_idx = -1
    for i, line in enumerate(lines):
        low = line.lower()
        norm = _normalize_match_text(line)
        if title_n and company_n and title_n in norm and company_n in norm:
            start_idx = i
            break
    if start_idx == -1 and title_n:
        for i, line in enumerate(lines):
            if title_n in _normalize_match_text(line):
                start_idx = i
                break
    if start_idx == -1 and company_n:
        for i, line in enumerate(lines):
            if company_n in _normalize_match_text(line):
                start_idx = i
                break
    if start_idx == -1:
        return []
    bullets = []
    for j in range(start_idx + 1, len(lines)):
        line = lines[j].strip()
        if not line:
            if bullets:
                break
            continue
        if re.match(r"^[-•*]\s+", line):
            bullets.append(re.sub(r"^[-•*]\s+", "", line).strip())
            continue
        if re.match(r"^\d+[\.\)]\s+", line):
            bullets.append(re.sub(r"^\d+[\.\)]\s+", "", line).strip())
            continue
        if line.lower().startswith("responsibilities:"):
            remainder = line.split(":", 1)[1].strip()
            if remainder:
                bullets.append(remainder)
            continue
        if re.match(r"^[A-Za-z].*\\d{4}", line):
            break
        if line.lower().startswith(("professional experience", "education", "certifications", "core skills", "skills", "professional summary")):
            break
    return bullets


def _build_experience_section(consultant, source_content=None, bullets_map=None, override_title=None):
    lines = ["PROFESSIONAL EXPERIENCE"]
    experiences = list(consultant.experience.all())
    if not experiences:
        lines.append("No experience listed.")
        return "\n".join(lines)
    bullets_map = bullets_map or {}
    ordered = [e for e, _, _ in _target_counts_for_experiences(experiences)]
    for idx, e in enumerate(ordered):
        start = _format_month_year(e.start_date)
        end = "Present" if e.is_current else _format_month_year(e.end_date)
        role_title = override_title if override_title and idx == 0 else e.title
        lines.append(f"{role_title} | {e.company} | {start} – {end}".strip())
        if e.description:
            for item in [x.strip() for x in e.description.splitlines() if x.strip()]:
                lines.append(f"- {item}")
        else:
            key = f"{_normalize_match_text(e.title)}||{_normalize_match_text(e.company)}"
            bullets = bullets_map.get(key, [])
            if not bullets and source_content:
                bullets = _extract_bullets_for_role(source_content, e.title, e.company)
            if not bullets:
                logger.warning("No bullets found for role: %s @ %s", e.title, e.company)
            for b in bullets:
                lines.append(f"- {b}")
        lines.append("")
    return "\n".join(lines).strip()


def _target_counts_for_experiences(experiences):
    items = list(experiences)
    if not items:
        return []
    def _sort_key(e):
        end_date = e.end_date or datetime.date.min
        start_date = e.start_date or datetime.date.min
        return (1 if e.is_current else 0, end_date, start_date)
    items_sorted = sorted(items, key=_sort_key, reverse=True)
    targets = []
    for idx, e in enumerate(items_sorted):
        if idx == 0:
            targets.append((e, 7, 10))
        else:
            targets.append((e, 6, 6))
    return targets


def generate_experience_bullets_with_counts(job, consultant, roles_needed, system_prompt=None):
    if not roles_needed:
        return {}
    llm = LLMService()
    if not llm.client:
        logger.warning("LLM client unavailable, cannot generate bullets.")
        return {}

    base_resume = consultant.base_resume_text or ""
    jd = _clean_jd_text(job.description or "")
    has_base = bool(base_resume.strip())

    if has_base:
        user_prompt = _fill_prompt(
            BULLETS_PROMPT_BASE,
            roles=json.dumps(roles_needed),
            jd=jd,
            base_resume=base_resume,
        )
    else:
        user_prompt = _fill_prompt(
            BULLETS_PROMPT_NO_BASE,
            roles=json.dumps(roles_needed),
            jd=jd,
        )
    system_prompt = system_prompt or BULLETS_SYSTEM_PROMPT
    content, _, error = llm.generate_with_prompts(job, consultant, system_prompt, user_prompt)
    if error or not content:
        logger.warning("Bullet generation failed: %s", error or "empty response")
        return {}
    logger.debug("Bullet generation raw response length: %s", len(content))
    try:
        data = json.loads(content)
    except Exception:
        match = re.search(r"\{.*\}", content, re.DOTALL)
        if not match:
            logger.warning("Bullet JSON parse failed, no JSON found")
            return {}
        try:
            data = json.loads(match.group(0))
        except Exception:
            logger.warning("Bullet JSON parse failed after extraction")
            return {}
    roles_out = data.get("roles") if isinstance(data, dict) else None
    if not roles_out:
        return {}
    bullets_map = {}
    for r in roles_out:
        title = (r.get("title") or "").strip()
        company = (r.get("company") or "").strip()
        bullets = [b.strip() for b in (r.get("bullets") or []) if str(b).strip()]
        if not title or not company or not bullets:
            continue
        key = f"{_normalize_match_text(title)}||{_normalize_match_text(company)}"
        bullets_map[key] = bullets
    # Second pass: enforce 22–25 words and structure if LLM ignored rules
    needs_fix = any(
        (_bullet_word_count(b) < 22 or _bullet_word_count(b) > 25)
        for bl in bullets_map.values()
        for b in bl
    )
    if needs_fix:
        repair_prompt = _fill_prompt(
            BULLETS_REPAIR_PROMPT,
            roles=json.dumps(roles_needed),
            bullets_map=json.dumps(bullets_map),
            jd=jd,
        )
        content2, _, error2 = llm.generate_with_prompts(job, consultant, system_prompt, repair_prompt)
        if not error2 and content2:
            try:
                data2 = json.loads(content2)
                roles2 = data2.get("roles") if isinstance(data2, dict) else None
                if roles2:
                    bullets_map = {}
                    for r in roles2:
                        title = (r.get("title") or "").strip()
                        company = (r.get("company") or "").strip()
                        bullets = [b.strip() for b in (r.get("bullets") or []) if str(b).strip()]
                        if not title or not company or not bullets:
                            continue
                        key = f"{_normalize_match_text(title)}||{_normalize_match_text(company)}"
                        bullets_map[key] = bullets
            except Exception:
                pass
    return bullets_map


def build_experience_bullets_map(job, consultant, source_content):
    bullets_map = {}
    needs = []
    targets = _target_counts_for_experiences(consultant.experience.all())
    method_keywords = _collect_method_keywords(job, consultant)
    required_terms = _required_terms_from_jd(job.description or "")
    for e, min_count, max_count in targets:
        base_bullets = []
        if e.description:
            for item in [x.strip() for x in e.description.splitlines() if x.strip()]:
                cleaned = _clean_bullet_line(item)
                if cleaned and not _is_skill_bullet(cleaned):
                    base_bullets.append(cleaned)
        else:
            base_bullets = _extract_bullets_for_role(source_content, e.title, e.company) if source_content else []
            base_bullets = [b for b in base_bullets if b and not _is_skill_bullet(b)]
        if max_count and len(base_bullets) > max_count:
            base_bullets = base_bullets[:max_count]
        # NOTE: Removed duplicate _expand_bullet_to_min_words + _apply_jd_alignment_rules
        # calls here — they already run in the final pass below (lines 878-882).
        base_bullets = _dedupe_bullets(base_bullets)
        key = f"{_normalize_match_text(e.title)}||{_normalize_match_text(e.company)}"
        bullets_map[key] = base_bullets
        if len(base_bullets) < min_count:
            needs.append({
                "title": e.title,
                "company": e.company,
                "count": min_count - len(base_bullets),
            })

    if needs:
        logger.warning("Missing experience bullets detected (enforce counts). roles=%s", [f"{n['title']} @ {n['company']}" for n in needs])
        generated_map = generate_experience_bullets_with_counts(job, consultant, needs)
        for n in needs:
            key = f"{_normalize_match_text(n['title'])}||{_normalize_match_text(n['company'])}"
            existing = bullets_map.get(key, [])
            generated = generated_map.get(key, [])
            if generated:
                for b in generated:
                    if b not in existing:
                        existing.append(b)
            bullets_map[key] = existing

    # Final trim to max counts
    for idx, (e, min_count, max_count) in enumerate(targets):
        key = f"{_normalize_match_text(e.title)}||{_normalize_match_text(e.company)}"
        items = bullets_map.get(key, [])
        if len(items) < min_count:
            logger.warning("LLM unavailable and missing bullets for role: %s", key)
        if max_count and len(items) > max_count:
            bullets_map[key] = items[:max_count]
        else:
            bullets_map[key] = items
        bullets_map[key] = [_fix_broken_bullet_text(b) for b in bullets_map[key]]
        bullets_map[key] = [_expand_bullet_to_min_words_strict(b, job, 22, method_keywords) for b in bullets_map[key]]
        bullets_map[key] = [_cap_bullet_words(b, 25) for b in bullets_map[key]]
        bullets_map[key] = _apply_metric_rules(bullets_map[key], job, method_keywords, _cap_metrics_for_role(job, idx))
        bullets_map[key] = _apply_jd_alignment_rules(bullets_map[key], job.description or "")
        bullets_map[key] = [_strip_vague_adverbs(b) for b in bullets_map[key]]
        bullets_map[key] = [_expand_bullet_to_min_words_strict(b, job, 22, method_keywords) for b in bullets_map[key]]
        bullets_map[key] = [_cap_bullet_words(b, 25) for b in bullets_map[key]]
        bullets_map[key] = [_fix_broken_bullet_text(b) for b in bullets_map[key]]
        bullets_map[key] = _dedupe_bullets(bullets_map[key])

    # Enforce required terms across experience bullets (auto-rewrite if missing)
    if required_terms:
        all_text = " ".join(
            " ".join(bullets_map.get(f"{_normalize_match_text(e.title)}||{_normalize_match_text(e.company)}", []))
            for e, _, _ in targets
        ).lower()
        missing = [t for t in required_terms if t not in all_text]
        if missing:
            llm = LLMService()
            if llm.client:
                roles_payload = []
                for e, _, _ in targets:
                    key = f"{_normalize_match_text(e.title)}||{_normalize_match_text(e.company)}"
                    roles_payload.append({
                        "title": e.title,
                        "company": e.company,
                        "bullets": bullets_map.get(key, [])
                    })
                rewrite_prompt = (
                    _fill_prompt(
                        BULLETS_MISSING_TERMS_PROMPT,
                        missing_terms=", ".join(missing),
                        roles_payload=json.dumps(roles_payload),
                    )
                )
                content, _, error = llm.generate_with_prompts(
                    job, consultant, BULLETS_SYSTEM_PROMPT, rewrite_prompt, force_new=True
                )
                if not error and content:
                    try:
                        data = json.loads(content)
                        roles_out = data.get("roles") if isinstance(data, dict) else None
                        if roles_out:
                            for r in roles_out:
                                title = (r.get("title") or "").strip()
                                company = (r.get("company") or "").strip()
                                bullets = [b.strip() for b in (r.get("bullets") or []) if str(b).strip()]
                                if not title or not company or not bullets:
                                    continue
                                key = f"{_normalize_match_text(title)}||{_normalize_match_text(company)}"
                                bullets_map[key] = bullets
                    except Exception:
                        pass

    # If JD-critical operational terms are still missing, enforce targeted bullets for the most recent role.
    jd_text = job.description or ""
    crit_terms = ["support tickets", "on-call", "on call", "xml", "rest api", "http", "response codes", "firewall", "network connectivity"]
    crit_missing = [t for t in crit_terms if t in jd_text.lower() and t not in all_text]
    if crit_missing and targets:
        llm = LLMService()
        if llm.client:
            # Most recent role is first in targets
            e, _, _ = targets[0]
            key = f"{_normalize_match_text(e.title)}||{_normalize_match_text(e.company)}"
            existing = bullets_map.get(key, [])
            prompt = _fill_prompt(
                BULLETS_MISSING_TERMS_PROMPT,
                missing_terms=", ".join(crit_missing),
                roles_payload=json.dumps(
                    [
                        {
                            "title": e.title,
                            "company": e.company,
                            "bullets": existing,
                        }
                    ]
                ),
            )
            content3, _, error3 = llm.generate_with_prompts(
                job, consultant, BULLETS_SYSTEM_PROMPT, prompt, force_new=True
            )
            if not error3 and content3:
                try:
                    data3 = json.loads(content3)
                    roles3 = data3.get("roles") if isinstance(data3, dict) else None
                    if roles3:
                        for r in roles3:
                            title = (r.get("title") or "").strip()
                            company = (r.get("company") or "").strip()
                            bullets = [b.strip() for b in (r.get("bullets") or []) if str(b).strip()]
                            if not title or not company or not bullets:
                                continue
                            key = f"{_normalize_match_text(title)}||{_normalize_match_text(company)}"
                            bullets_map[key] = bullets
                except Exception:
                    pass
    return bullets_map



def _build_education_section(consultant):
    lines = ["EDUCATION"]
    educations = consultant.education.all()
    if not educations:
        lines.append("No education listed.")
        return "\n".join(lines)
    for e in educations:
        end = _format_month_year(e.end_date) if e.end_date else "Present"
        program = e.degree or "Degree"
        if e.field_of_study:
            program = f"{program} in {e.field_of_study}"
        lines.append(f"{program} | {e.institution} | {end}")
    return "\n".join(lines)


def _build_certifications_section(consultant):
    certs = consultant.certifications.all()
    if not certs:
        return ""
    lines = ["CERTIFICATIONS"]
    for c in certs:
        lines.append(f"- {c.name}")
    return "\n".join(lines)


def _build_skills_section(job, consultant):
    required_terms = _required_terms_from_jd(job.description or "")
    return generate_skills_from_jd(job, required_terms, consultant)


def _extract_metrics_from_text(text):
    if not text:
        return []
    metrics = re.findall(r"(\\b\\d+\\s*%|\\b\\d+\\s*(?:ms|s|sec|secs|seconds|min|mins|minutes|hours|hrs|days|x)\\b|\\b\\d+\\s*(?:%|percent)\\b)", text, flags=re.I)
    # Deduplicate while preserving order
    seen = set()
    out = []
    for m in metrics:
        if m.lower() in seen:
            continue
        seen.add(m.lower())
        out.append(m)
    return out


def _normalize_title(t):
    if not t:
        return ""
    t = t.lower()
    t = re.sub(r"\b(senior|sr|jr|junior|lead|principal|i|ii|iii|iv|v)\b", "", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def _term_present(summary_lower, term):
    if not term:
        return False
    t = term.lower()
    if t in summary_lower:
        return True
    # For multi-word terms, allow token presence
    tokens = [tok for tok in re.findall(r"[A-Za-z0-9']+", t) if tok]
    if len(tokens) >= 2:
        return all(tok in summary_lower for tok in tokens)
    return False


def _sanitize_summary(summary, title, years_display):
    if not summary:
        return summary
    text = " ".join(summary.strip().split())
    # Drop generic qualifier phrases in summary to keep it technical and concrete.
    text = re.sub(r"\b(basic understanding of|knowledge of|familiarity with|general)\b\s+", "", text, flags=re.I)
    # Remove sentences containing JD meta or work-arrangement language.
    sentences = re.split(r"(?<=[.!?])\s+", text)
    filtered = []
    for s in sentences:
        if re.search(r"\b(hybrid|onsite|on-site|remote|work arrangement|arrangements|vision)\b", s, re.I):
            continue
        filtered.append(s)
    text = " ".join(filtered).strip()
    # Ensure summary starts with the exact job title.
    if title and not text.lower().startswith(title.lower()):
        if years_display and f"{years_display} years" in text:
            remainder = re.sub(r"^[^ ]+.*?\bwith\b\s+", "", text, flags=re.I)
        else:
            remainder = text
        prefix = f"{title} with {years_display} years of experience "
        text = prefix + remainder
    # Ensure years appear
    if years_display and f"{years_display} years" not in text:
        text = f"{title} with {years_display} years of experience " + text[len(title):].lstrip()
    return " ".join(text.split())


def _validate_summary(summary, title, years_display, jd_keywords, required_terms=None):
    reasons = []
    if not summary:
        reasons.append("empty")
        return False, reasons
    lines = [l for l in summary.splitlines() if l.strip()]
    if len(lines) < 4 or len(lines) > 6:
        reasons.append(f"line_count={len(lines)}")
    if any(l.strip().startswith(("-", "•", "*")) for l in lines):
        reasons.append("bullet_line_found")
    # No pronouns
    if re.search(r"\b(i|my|me|we|our|us|he|she|his|her)\b", summary, re.I):
        reasons.append("pronoun_found")
    # No buzzwords / generic phrases
    if re.search(r"\b(dynamic|passionate|ninja|rockstar|guru|world-class|go-getter)\b", summary, re.I):
        reasons.append("buzzword_found")
    if re.search(r"\b(proven track record|innovative solutions|results-driven|cutting-edge|fast-paced)\b", summary, re.I):
        reasons.append("generic_phrase_found")
    if re.search(r"\b(minimum|maximum|salary|compensation|benefits|posting date)\b", summary, re.I):
        reasons.append("jd_noise_terms")
    if re.search(r"\b(hybrid|onsite|on-site|remote|work arrangement|arrangements)\b", summary, re.I):
        reasons.append("work_arrangement_term")
    if re.search(r"\b(vision for|contributing to the vision|vision)\b", summary, re.I):
        reasons.append("vision_phrase")
    # Must include title and years on line 1
    if not lines or not lines[0].lower().startswith(title.lower()):
        # allow normalized title match (e.g., DevOps Engineer I -> DevOps Engineer)
        norm_title = _normalize_title(title)
        if norm_title and not lines[0].lower().startswith(norm_title):
            reasons.append("title_missing")
    if years_display and f"{years_display} years".lower() not in summary.lower():
        reasons.append("years_missing")
    # Must include 4+ JD keywords (exact tokens)
    hits = 0
    for k in jd_keywords:
        if k in summary.lower():
            hits += 1
    if hits < 4:
        reasons.append("jd_keywords<4")
    # Must include required terms when present
    required_terms = [t for t in (required_terms or []) if t]
    summary_lower = summary.lower()
    missing = [t for t in required_terms if not _term_present(summary_lower, t)]
    if missing:
        reasons.append("required_terms_missing")
    return (len(reasons) == 0), reasons


def _pick_top_keywords(jd_text, limit=6):
    keywords = extract_keywords(jd_text or "", max_keywords=200)
    # Prefer longer/technical tokens over generic ones and remove numerics
    keywords = [k for k in keywords if len(k) >= 3 and not k.isdigit()]
    return keywords[:limit]


def _build_summary_section(job, consultant):
    jd_text = job.description or ""
    title = job.title or "DevOps Engineer"
    years_display = _total_experience_years_display(consultant) or "0+"

    parsed = getattr(job, "parsed_jd", None) or {}
    if not parsed:
        return "PROFESSIONAL SUMMARY\nSummary unavailable (JD not parsed)."

    jd_keywords = [k.lower() for k in (parsed.get("keywords_for_ats") or []) if isinstance(k, str)]
    required_terms_summary = [k.lower() for k in (parsed.get("required_skills") or []) if isinstance(k, str)]
    required_terms_summary = required_terms_summary[:4]
    metrics = []
    for exp in consultant.experience.all():
        metrics.extend(_extract_metrics_from_text(exp.description or ""))

    exp_summary = []
    for exp in consultant.experience.all():
        exp_summary.append(f"{exp.title} at {exp.company} ({exp.start_date.strftime('%b %Y') if exp.start_date else ''}–{'Present' if exp.is_current else (exp.end_date.strftime('%b %Y') if exp.end_date else '')})")
    exp_summary_text = "; ".join([e for e in exp_summary if e]) or "No experience listed"

    llm = LLMService()
    last_fail_reasons = []
    if llm.client:
        system_prompt = (
            "You are a resume writer. Return ONLY the summary text as 4–6 lines, no bullets, no labels."
        )
        user_prompt = (
            "STEP 2 — WRITE THE PROFESSIONAL SUMMARY USING THIS FORMULA\n"
            "Line 1: [Exact Job Title] with [X] years of experience [top 2 responsibilities using JD's exact language]\n"
            "Line 2: Proven expertise in [4–5 exact keywords from the JD] with demonstrated impact in [primary outcome the JD cares about]\n"
            "Line 3: [A strength that mirrors the JD's requirements section — use the JD's own phrasing]\n"
            "Line 4: [A line that mirrors the JD's culture tone — collaborative, mission-critical, data-driven, etc.]\n"
            "Lines 5–6: optional only if needed to include required JD wording.\n"
            "\nSTRICT RULES:\n"
            "- Return ONLY the summary lines. No labels, no explanations, no bullet points.\n"
            "- Exactly 4–6 lines (line breaks required).\n"
            "- Use exact keywords from the JD analysis.\n"
            "- Do not use: \"dynamic professional\", \"results-oriented\", \"passionate about\", or any first person.\n"
            "- The summary must feel like the JD wrote it about the candidate.\n"
            f"- Use this exact job title: {title}\n"
            f"- Use this exact years format: {years_display} years\n"
            "\nPARSED JD JSON:\n"
            f"{json.dumps(parsed, ensure_ascii=True)}\n"
        )
        content, _, error = llm.generate_with_prompts(job, consultant, system_prompt, user_prompt)
        candidate = ""
        if not error and content:
            candidate = "\n".join([l.strip() for l in content.strip().splitlines() if l.strip()])
            ok, reasons = _validate_summary(candidate, title, years_display, jd_keywords, required_terms_summary)
            last_fail_reasons = reasons
            if ok:
                return "PROFESSIONAL SUMMARY\n" + candidate
            logger.warning("Summary candidate rejected: %s", candidate)

        # Second attempt with tighter instructions
        if required_terms_summary:
            user_prompt += "\\nInclude these exact terms: " + ", ".join(required_terms_summary)
        user_prompt += "\\nRewrite to meet all rules exactly."
        content2, _, error2 = llm.generate_with_prompts(job, consultant, system_prompt, user_prompt)
        if not error2 and content2:
            candidate = "\n".join([l.strip() for l in content2.strip().splitlines() if l.strip()])
            ok, reasons = _validate_summary(candidate, title, years_display, jd_keywords, required_terms_summary)
            last_fail_reasons = reasons
            if ok:
                return "PROFESSIONAL SUMMARY\n" + candidate
            logger.warning("Summary candidate rejected (retry): %s", candidate)
        if last_fail_reasons:
            logger.warning("Summary validation failed: %s", ", ".join(last_fail_reasons))
        if candidate:
            return "PROFESSIONAL SUMMARY\n" + candidate

    # If LLM unavailable, return a clear note.
    return "PROFESSIONAL SUMMARY\nSummary unavailable (LLM not configured)."


def normalize_generated_resume(content, job, consultant, bullets_map=None):
    header_block = _build_header_block(job, consultant)
    summary_section = _build_summary_section(job, consultant)
    skills_section = _build_skills_section(job, consultant)
    edu_section = _build_education_section(consultant)
    cert_section = _build_certifications_section(consultant)

    if bullets_map is None:
        bullets_map = build_experience_bullets_map(job, consultant, "")
    # Decide whether to match most recent role title to the JD title
    match_override = consultant.match_jd_title_override
    if match_override is None:
        match_override = PlatformConfig.load().match_jd_title_default
    override_title = job.title if match_override else None

    exp_section = _build_experience_section(
        consultant,
        source_content="",
        bullets_map=bullets_map,
        override_title=override_title
    )

    # Assemble a clean, deterministic resume using the required structure.
    parts = [
        header_block,
        "",
        summary_section,
        "",
        skills_section,
        "",
        exp_section,
        "",
        edu_section,
    ]
    if cert_section:
        parts += ["", cert_section]

    text = "\n".join(parts).strip()

    text = _normalize_core_skills_format(text)

    return text.strip()
SECTION_KEYS = {
    "name",
    "email",
    "phone",
    "jd_location",
    "professional_summary",
    "skills",
    "base_resume",
    "experience",
    "education",
    "jd_description",
}


def build_user_prompt_from_sections(job, consultant, sections):
    selected = set(sections or [])
    selected = selected.intersection(SECTION_KEYS)

    parts = [
        "STRICT DATA RULES:",
        "- Use the provided profile data exactly for name, contact, experience titles/companies/dates, education, and certifications.",
        "- Do NOT invent or replace people, companies, dates, degrees, or certifications.",
        "- If certifications are not provided, do NOT add a Certifications section.",
        "- Bullet counts: most recent role must have 7–10 bullets; all other roles must have exactly 6 bullets.",
        "- Do NOT repeat the same sentence or phrase across bullets or roles.",
        "- Latest role title MUST exactly match the JD role title.",
        "- Every bullet must be 22–25 words.",
        "- Naturally integrate JD-required keywords (e.g., Windows, Linux, IIS, REST, XML, Firewall) into sentences. Do NOT just list them.",
        "- Each bullet must follow: Action + Tool/Method + Context + Outcome.",
        "- Example: 'Engineered a scalable CI/CD pipeline using Jenkins and Docker, reducing deployment cycle times by 40% and ensuring 99.9% uptime.'",
        "- Output plain text only. No markdown headings, no bold, no tables.",
        "- Use this structure exactly:",
        "  FULL NAME",
        "  City, State | Email | Phone | LinkedIn (only if provided)",
        "  PROFESSIONAL SUMMARY",
        "  SKILLS",
        "  PROFESSIONAL EXPERIENCE",
        "  EDUCATION",
        "  CERTIFICATIONS (only if provided)",
        "",
    ]
    contact_name = consultant.user.get_full_name() or consultant.user.username
    contact_email = consultant.user.email or "Not provided."
    contact_phone = consultant.phone or "Not provided."

    if "name" in selected:
        parts.append(f"Name: {contact_name}")
    if "email" in selected:
        parts.append(f"Email: {contact_email}")
    if "phone" in selected:
        parts.append(f"Phone: {contact_phone}")
    if "jd_location" in selected:
        parts.append(f"Location (use JD location): {job.location or 'Not provided.'}")

    years_display = _total_experience_years_display(consultant)
    if years_display:
        parts.append(f"Total Experience (use exactly): {years_display} years")

    if "professional_summary" in selected:
        parts.append(BUILD_PROMPT_SUMMARY_INSTRUCTION)

    if "skills" in selected:
        skills = consultant.skills or []
        if skills:
            parts.append(BUILD_PROMPT_SKILLS_WITH_PROFILE)
            parts.append(f"Skills (from profile): {', '.join(skills)}")
        else:
            parts.append(BUILD_PROMPT_SKILLS_GENERATE)

    if "base_resume" in selected:
        base_resume_text = consultant.base_resume_text or ""
        if base_resume_text.strip():
            parts.append(BUILD_PROMPT_BASE_RESUME_LABEL)
            parts.append(base_resume_text)
        else:
            parts.append(BUILD_PROMPT_BASE_RESUME_MISSING)
            parts.append(BUILD_PROMPT_BASE_RESUME_GUIDE_1)
            parts.append(BUILD_PROMPT_BASE_RESUME_GUIDE_2)
            parts.append(BUILD_PROMPT_BASE_RESUME_GUIDE_3)

    if "experience" in selected:
        parts.append(BUILD_PROMPT_EXPERIENCE_LABEL)
        experiences = list(consultant.experience.all())
        ordered = [e for e, _, _ in _target_counts_for_experiences(experiences)]
        if experiences:
            for e in ordered:
                start = e.start_date.strftime('%Y') if e.start_date else ''
                end = "Present" if e.is_current else (e.end_date.strftime('%Y') if e.end_date else '')
                role_line = f"{e.title} | {e.company} | {start} – {end}"
                if e.description:
                    role_line += f"\n  Responsibilities: {e.description}"
                else:
                    role_line += f"\n  {BUILD_PROMPT_ROLE_NO_DESC}"
                parts.append(role_line)
        else:
            parts.append(BUILD_PROMPT_NO_EXPERIENCE)

    if "education" in selected:
        parts.append(BUILD_PROMPT_EDUCATION_LABEL)
        educations = consultant.education.all()
        if educations:
            for e in educations:
                start = e.start_date.strftime('%Y') if e.start_date else ''
                end = e.end_date.strftime('%Y') if e.end_date else 'Present'
                year = end if end else start
                parts.append(f"{e.degree} | {e.institution} | {year}")
        else:
            parts.append(BUILD_PROMPT_NO_EDUCATION)

    if "jd_description" in selected:
        parts.append(BUILD_PROMPT_JD_HEADER)
        parts.append(job.description or "Not provided.")

    return "\n".join(parts).strip()

def get_system_prompt_text(job, consultant, prompt_override=None):
    if prompt_override:
        if prompt_override.system_text:
            return prompt_override.system_text
        if prompt_override.description:
            return strip_tags(prompt_override.description)
    # Force resume-specific prompt when available
    resume_prompt = Prompt.objects.filter(name='resume-2').first()
    if resume_prompt:
        if resume_prompt.system_text:
            return resume_prompt.system_text
        if resume_prompt.description:
            return strip_tags(resume_prompt.description)
    prompt = get_active_prompt_for_job(job, consultant)
    if prompt:
        if prompt.system_text:
            return prompt.system_text
        if prompt.description:
            return strip_tags(prompt.description)
    return DEFAULT_SYSTEM_PROMPT


def build_input_summary(job, consultant):
    experiences = []
    for exp in consultant.experience.all():
        experiences.append({
            'title': exp.title,
            'company': exp.company,
            'start_year': exp.start_date.strftime('%Y') if exp.start_date else '',
            'end_year': '' if exp.is_current or not exp.end_date else exp.end_date.strftime('%Y'),
            'is_current': exp.is_current,
        })

    educations = []
    for edu in consultant.education.all():
        educations.append({
            'degree': edu.degree,
            'field_of_study': edu.field_of_study,
            'institution': edu.institution,
            'start_year': edu.start_date.strftime('%Y') if edu.start_date else '',
            'end_year': edu.end_date.strftime('%Y') if edu.end_date else 'Present',
        })

    return {
        'job_title': job.title,
        'job_company': job.company,
        'job_location': job.location or 'Not provided.',
        'job_description': job.description,
        'consultant_name': consultant.user.get_full_name() or consultant.user.username,
        'consultant_email': consultant.user.email or 'Not provided.',
        'consultant_phone': consultant.phone or 'Not provided.',
        'base_resume_text': consultant.base_resume_text or '',
        'experience': experiences,
        'education': educations,
    }


class LLMService:

    def __init__(self):
        config = LLMConfig.load()
        self.config = config
        self.api_key = decrypt_value(config.encrypted_api_key)
        if self.api_key and not self.api_key.startswith('sk-your') and config.generation_enabled:
            self.client = openai.OpenAI(api_key=self.api_key)
        else:
            self.client = None

    def _build_prompt(self, job, consultant, prompt_override=None):
        """Build the user prompt from template or default."""
        # Gather contact info
        contact_name = consultant.user.get_full_name() or consultant.user.username
        contact_email = consultant.user.email or "Not provided."
        contact_phone = consultant.phone or "Not provided."
        base_resume_text = consultant.base_resume_text or ""

        # Gather experience summary
        experiences = consultant.experience.all()
        exp_summary = "\n".join(
            f"- {e.title} at {e.company} ({e.start_date.strftime('%Y')}–{'Present' if e.is_current else e.end_date.strftime('%Y') if e.end_date else ''})"
            for e in experiences
        ) or "No experience listed."

        # Gather education summary
        educations = consultant.education.all()
        edu_summary = "\n".join(
            f"- {e.degree} in {e.field_of_study} at {e.institution} ({e.start_date.strftime('%Y')}–{e.end_date.strftime('%Y') if e.end_date else 'Present'})"
            for e in educations
        ) or "No education listed."

        # Gather certifications
        certs = consultant.certifications.all()
        cert_summary = ", ".join(c.name for c in certs) or "None listed."

        input_summary = (
            f"Job: {job.title} @ {job.company}\n"
            f"Job Location: {job.location or 'Not provided.'}\n"
            f"Consultant: {contact_name}\n"
            f"Email: {contact_email}\n"
            f"Phone: {contact_phone}\n"
            f"Experience:\n{exp_summary}\n"
            f"Education:\n{edu_summary}\n"
            f"Certifications: {cert_summary}\n"
        )

        prompt = prompt_override or get_active_prompt_for_job(job, consultant)
        if prompt:
            try:
                template_text = prompt.template_text
                base = template_text.format(
                    job_title=job.title,
                    company=job.company,
                    job_description=job.description,
                    consultant_name=consultant.user.get_full_name() or consultant.user.username,
                    consultant_bio=consultant.bio or "Not provided.",
                    consultant_skills=", ".join(consultant.skills) if consultant.skills else "Not provided.",
                    experience_summary=exp_summary,
                    certifications=cert_summary,
                    base_resume_text=base_resume_text,
                    input_summary=input_summary,
                )
                return (
                    f"{base}\n\n"
                    f"{BUILD_PROMPT_TEMPLATE_SUFFIX.format(jd_text=job.description or 'Not provided.')}"
                )
            except (KeyError, IndexError):
                pass  # Fall through to default

        base_section = (
            BUILD_PROMPT_BASE_SECTION_WITH.format(base_resume_text=base_resume_text)
            if base_resume_text.strip()
            else BUILD_PROMPT_BASE_SECTION_WITHOUT
        )

        return (
            f"Consultant Name: {contact_name}\n"
            f"Consultant Email: {contact_email}\n"
            f"Consultant Phone: {contact_phone}\n"
            f"Bio: {consultant.bio or 'Not provided.'}\n"
            f"Skills: {', '.join(consultant.skills) if consultant.skills else 'Not provided.'}\n"
            f"{base_section}"
            f"Experience:\n{exp_summary}\n"
            f"Education:\n{edu_summary}\n"
            f"Certifications: {cert_summary}\n\n"
            f"--- TARGET JOB ---\n"
            f"Title: {job.title}\n"
            f"Company: {job.company}\n"
            f"Description:\n{job.description}\n"
            f"\n{BUILD_PROMPT_REQUIRED_SECTIONS}"
        )

    def generate_resume_content(self, job, consultant, actor=None, prompt_override=None, force_new=False):
        """Generate resume content. Returns (content, tokens_used, error)."""
        prompt_text = self._build_prompt(job, consultant, prompt_override=prompt_override)
        system_prompt = get_system_prompt_text(job, consultant, prompt_override=prompt_override)

        if not self.client:
            mock = (
                f"PROFESSIONAL SUMMARY\n"
                f"Results-driven professional with expertise in "
                f"{', '.join(consultant.skills[:3]) if consultant.skills else 'various technologies'}. "
                f"Seeking the {job.title} position at {job.company}.\n\n"
                f"SKILLS\n"
            )
            if consultant.skills:
                for skill in consultant.skills:
                    mock += f"- {skill}\n"
            else:
                mock += "- Skills not listed\n"

            mock += (
                f"\nPROFESSIONAL EXPERIENCE\n"
                f"(Experience details from profile)\n\n"
                f"EDUCATION\n"
                f"(Education details from profile)\n"
            )
            return mock, 0, None

        if self.config.monthly_token_cap:
            month_start = timezone.now().replace(day=1, hour=0, minute=0, second=0, microsecond=0)
            total_month_tokens = LLMUsageLog.objects.filter(created_at__gte=month_start).aggregate(
                total=Sum('total_tokens')
            )['total'] or 0
            if total_month_tokens >= self.config.monthly_token_cap and self.config.auto_disable_on_cap:
                self.config.generation_enabled = False
                self.config.save()
                return None, 0, "Monthly token cap reached. Generation disabled."

        request_payload = {
            "model": self.config.active_model or "gpt-4o-mini",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt_text},
            ],
            "temperature": float(self.config.temperature),
            "max_tokens": self.config.max_output_tokens,
        }

        try:
            start = time.time()
            response = self.client.chat.completions.create(
                model=request_payload["model"],
                messages=request_payload["messages"],
                temperature=request_payload["temperature"],
                max_tokens=request_payload["max_tokens"],
            )
            latency_ms = int((time.time() - start) * 1000)
            content = response.choices[0].message.content
            prompt_tokens = response.usage.prompt_tokens if response.usage else 0
            completion_tokens = response.usage.completion_tokens if response.usage else 0
            tokens = response.usage.total_tokens if response.usage else 0
            costs = calculate_cost(self.config.active_model, prompt_tokens, completion_tokens)
            LLMUsageLog.objects.create(
                model_name=self.config.active_model,
                system_prompt=system_prompt,
                user_prompt=prompt_text,
                request_payload=request_payload,
                response_text=content or "",
                prompt_tokens=prompt_tokens,
                completion_tokens=completion_tokens,
                total_tokens=tokens,
                cost_input=costs['input'],
                cost_output=costs['output'],
                cost_total=costs['total'],
                latency_ms=latency_ms,
                success=True,
                job=job,
                consultant=consultant,
                actor=actor,
            )
            return content, tokens, None
        except Exception as e:
            LLMUsageLog.objects.create(
                model_name=self.config.active_model,
                success=False,
                error_message=str(e),
                system_prompt=system_prompt,
                user_prompt=prompt_text,
                request_payload=request_payload,
                job=job,
                consultant=consultant,
                actor=actor,
            )
            return None, 0, str(e)

    def generate_with_prompts(
        self,
        job,
        consultant,
        system_prompt,
        user_prompt,
        actor=None,
        force_new=False,
        temperature_override=None,
    ):
        """Generate resume content using explicit prompts. Returns (content, tokens_used, error)."""
        if not self.client:
            mock = (
                f"PROFESSIONAL SUMMARY\n"
                f"Results-driven professional with expertise in "
                f"{', '.join(consultant.skills[:3]) if consultant.skills else 'various technologies'}. "
                f"Seeking the {job.title} position at {job.company}.\n\n"
                f"SKILLS\n"
            )
            if consultant.skills:
                for skill in consultant.skills:
                    mock += f"- {skill}\n"
            else:
                mock += "- Skills not listed\n"

            mock += (
                f"\nPROFESSIONAL EXPERIENCE\n"
                f"(Experience details from profile)\n\n"
                f"EDUCATION\n"
                f"(Education details from profile)\n"
            )
            return mock, 0, None

        if self.config.monthly_token_cap:
            month_start = timezone.now().replace(day=1, hour=0, minute=0, second=0, microsecond=0)
            total_month_tokens = LLMUsageLog.objects.filter(created_at__gte=month_start).aggregate(
                total=Sum('total_tokens')
            )['total'] or 0
            if total_month_tokens >= self.config.monthly_token_cap and self.config.auto_disable_on_cap:
                self.config.generation_enabled = False
                self.config.save()
                return None, 0, "Monthly token cap reached. Generation disabled."

        temperature = float(self.config.temperature)
        if temperature_override is not None:
            temperature = float(temperature_override)

        request_payload = {
            "model": self.config.active_model or "gpt-4o-mini",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "temperature": temperature,
            "max_tokens": self.config.max_output_tokens,
        }

        try:
            start = time.time()
            response = self.client.chat.completions.create(
                model=request_payload["model"],
                messages=request_payload["messages"],
                temperature=request_payload["temperature"],
                max_tokens=request_payload["max_tokens"],
            )
            latency_ms = int((time.time() - start) * 1000)
            content = response.choices[0].message.content
            prompt_tokens = response.usage.prompt_tokens if response.usage else 0
            completion_tokens = response.usage.completion_tokens if response.usage else 0
            tokens = response.usage.total_tokens if response.usage else 0
            costs = calculate_cost(self.config.active_model, prompt_tokens, completion_tokens)
            LLMUsageLog.objects.create(
                model_name=self.config.active_model,
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                request_payload=request_payload,
                response_text=content or "",
                prompt_tokens=prompt_tokens,
                completion_tokens=completion_tokens,
                total_tokens=tokens,
                cost_input=costs['input'],
                cost_output=costs['output'],
                cost_total=costs['total'],
                latency_ms=latency_ms,
                success=True,
                job=job,
                consultant=consultant,
                actor=actor,
            )
            return content, tokens, None
        except Exception as e:
            LLMUsageLog.objects.create(
                model_name=self.config.active_model,
                success=False,
                error_message=str(e),
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                request_payload=request_payload,
                job=job,
                consultant=consultant,
                actor=actor,
            )
            return None, 0, str(e)


class DocxService:
    FONT_NAME = "Calibri"
    FONT_SIZE_BODY = Pt(11)
    FONT_SIZE_H1 = Pt(12)
    FONT_SIZE_H2 = Pt(12)
    FONT_SIZE_NAME = Pt(16)
    FONT_SIZE_CONTACT = Pt(11)

    def _set_font(self, run, size=None, bold=False):
        """Apply Aptos font to a run."""
        run.font.name = self.FONT_NAME
        if size: 
            run.font.size = size
        run.font.bold = bold

    def _add_formatted_paragraph(self, doc, text, style=None, alignment=None, font_size=None, bold=False, space_after=None, space_before=None):
        """Add a paragraph with Aptos font and optional formatting."""
        p = doc.add_paragraph(style=style)
        if alignment is not None:
            p.alignment = alignment
        if space_after is not None:
            p.paragraph_format.space_after = space_after
        if space_before is not None:
            p.paragraph_format.space_before = space_before
        p.paragraph_format.line_spacing = 1.0

        # Parse markdown bold (**text**) into actual bold runs
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                self._set_font(run, size=font_size or self.FONT_SIZE_BODY, bold=True)
            elif part:
                run = p.add_run(part)
                self._set_font(run, size=font_size or self.FONT_SIZE_BODY, bold=bold)
        return p

    def _add_thin_rule(self, doc):
        """Add a thin horizontal line separator."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # Use a border-bottom on the paragraph
        from docx.oxml.ns import qn
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def create_docx(self, content):
        """Convert markdown-ish text content into a properly formatted DOCX document."""
        doc = Document()

        # Set default font for the document
        style = doc.styles['Normal']
        style.font.name = self.FONT_NAME
        style.font.size = self.FONT_SIZE_BODY

        # Set margins (0.7 inch for a tight 2-page resume)
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        lines = content.split('\n')
        is_first_line = True
        pending_contact_line = False
        heading_set = {
            "PROFESSIONAL SUMMARY",
            "SKILLS",
            "PROFESSIONAL EXPERIENCE",
            "EDUCATION",
            "CERTIFICATIONS",
        }
        current_section = None

        for i, line in enumerate(lines):
            stripped = line.strip()

            # Skip empty lines
            if not stripped:
                continue

            # Horizontal rule (___) → thin line separator
            if stripped in ('___', '---', '***', '_ _ _'):
                self._add_thin_rule(doc)
                continue

            # H1 heading (# Section)
            if stripped.startswith('# ') and not stripped.startswith('## '):
                heading_text = stripped[2:].strip()
                self._add_formatted_paragraph(
                    doc, heading_text, bold=True,
                    font_size=self.FONT_SIZE_H1,
                    space_before=Pt(6), space_after=Pt(3)
                )
                continue

            # H2 heading (## Section)
            if stripped.startswith('## '):
                heading_text = stripped[3:].strip()
                if not heading_text:
                    continue  # Skip empty ## lines
                self._add_formatted_paragraph(
                    doc, heading_text, bold=True,
                    font_size=self.FONT_SIZE_H2,
                    space_before=Pt(6), space_after=Pt(3)
                )
                continue

            # Name/Header lines (first line = name, second line with pipes = contact)
            if is_first_line:
                is_first_line = False
                if '|' in stripped:
                    parts = [p.strip() for p in stripped.split('|')]
                    name = parts[0] if parts else stripped
                    self._add_formatted_paragraph(
                        doc, name, bold=True,
                        font_size=self.FONT_SIZE_NAME,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_after=Pt(1)
                    )
                    if len(parts) > 1:
                        contact = " | ".join(parts[1:])
                        self._add_formatted_paragraph(
                            doc, contact,
                            font_size=self.FONT_SIZE_CONTACT,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_after=Pt(4)
                        )
                    continue
                # Name only line
                self._add_formatted_paragraph(
                    doc, stripped, bold=True,
                    font_size=self.FONT_SIZE_NAME,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=Pt(1)
                )
                pending_contact_line = True
                continue

            if pending_contact_line and '|' in stripped:
                pending_contact_line = False
                self._add_formatted_paragraph(
                    doc, stripped,
                    font_size=self.FONT_SIZE_CONTACT,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=Pt(4)
                )
                continue

            is_first_line = False

            # Uppercase section heading
            if stripped in heading_set:
                self._add_formatted_paragraph(
                    doc, stripped, bold=True,
                    font_size=self.FONT_SIZE_H2,
                    space_before=Pt(6), space_after=Pt(4)
                )
                current_section = stripped
                continue

            # Experience role line (Title | Company | Dates) - bold title only
            if current_section == "PROFESSIONAL EXPERIENCE" and " | " in stripped and not stripped.startswith('- '):
                parts = [p.strip() for p in stripped.split('|')]
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.line_spacing = 1.0
                if parts:
                    run = p.add_run(parts[0])
                    self._set_font(run, size=self.FONT_SIZE_BODY, bold=True)
                    if len(parts) > 1:
                        rest = " | " + " | ".join(parts[1:])
                        run2 = p.add_run(rest)
                        self._set_font(run2, size=self.FONT_SIZE_BODY, bold=False)
                continue

            # Bullet point (- text)
            if stripped.startswith('- '):
                bullet_text = stripped[2:]
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.line_spacing = 1.0
                # Parse bold within bullet
                parts = re.split(r'(\*\*.*?\*\*)', bullet_text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        self._set_font(run, size=self.FONT_SIZE_BODY, bold=True)
                    elif part:
                        run = p.add_run(part)
                        self._set_font(run, size=self.FONT_SIZE_BODY)
                continue

            # Regular paragraph (with bold parsing)
            self._add_formatted_paragraph(doc, stripped, space_after=Pt(2))

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
from .skills_extractor import generate_skills_from_jd, extract_required_terms_from_jd

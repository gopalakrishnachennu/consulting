"""
Resume export utilities — DOCX and PDF.

DOCX: python-docx  (already in requirements)
PDF:  xhtml2pdf    (added to requirements)
"""
from __future__ import annotations

import re
from io import BytesIO
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    pass


# Default section order / labels / visibility. Each section can be reordered,
# hidden, or relabelled from the editor's Customise panel.
DEFAULT_SECTION_LAYOUT = [
    {'key': 'summary',        'label': 'Professional Summary',    'visible': True},
    {'key': 'skills',         'label': 'Technical Skills',        'visible': True},
    {'key': 'experience',     'label': 'Professional Experience', 'visible': True},
    {'key': 'education',      'label': 'Education',                'visible': True},
    {'key': 'certifications', 'label': 'Certifications',           'visible': True},
]

# Fallback template values — match render_resume_html's .get() defaults exactly,
# so a draft with no selected template still renders identically in preview + export.
DEFAULT_TEMPLATE = {
    'font_family': 'Georgia, serif',
    'name_size': 22, 'header_size': 13, 'body_size': 11, 'contact_size': 10,
    'accent_color': '#1e3a5f', 'name_color': '#111827', 'body_color': '#374151',
    'margin_top': 0.75, 'margin_bottom': 0.75, 'margin_left': 0.75, 'margin_right': 0.75,
    'line_height': 1.3, 'para_spacing': 5, 'section_spacing': 10,
    'header_style': 'underline', 'bullet_char': '•',
    # Phase 2 layout controls
    'sections_layout': DEFAULT_SECTION_LAYOUT,
    'skills_layout': 'categorized',  # 'categorized' | 'inline'
}


# ─── Colour helper ───────────────────────────────────────────────────────────

def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip('#')
    if len(h) == 3:
        h = ''.join(c * 2 for c in h)
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _esc(text: str) -> str:
    """HTML-escape a string."""
    return (
        str(text or '')
        .replace('&', '&amp;')
        .replace('<', '&lt;')
        .replace('>', '&gt;')
        .replace('"', '&quot;')
    )


_BOLD_RE = re.compile(r'\*\*(.+?)\*\*', re.S)


def _rich(text: str) -> str:
    """HTML-escape, then render **bold** markdown as <strong> for inline keyword emphasis."""
    return _BOLD_RE.sub(r'<strong>\1</strong>', _esc(text))


def _bold_segments(text: str):
    """Split text into (segment, is_bold) tuples on **bold** markers — for DOCX runs."""
    out, last, bold = [], 0, False
    for m in _BOLD_RE.finditer(str(text or '')):
        if m.start() > last:
            out.append((text[last:m.start()], False))
        out.append((m.group(1), True))
        last = m.end()
    if last < len(str(text or '')):
        out.append((str(text)[last:], False))
    return out or [(str(text or ''), False)]


# ─── HTML renderer (shared between preview & PDF) ────────────────────────────

def render_resume_html(sections: dict, tpl: dict, for_print: bool = False) -> str:
    """
    Render sections + template config to a standalone HTML string.
    When for_print=True, wraps in a full <html> document with @page CSS.
    """
    font       = tpl.get('font_family', 'Georgia, serif')
    name_sz    = tpl.get('name_size', 22)
    hdr_sz     = tpl.get('header_size', 13)
    body_sz    = tpl.get('body_size', 11)
    contact_sz = tpl.get('contact_size', 10)
    accent     = tpl.get('accent_color', '#1e3a5f')
    name_col   = tpl.get('name_color', '#111827')
    body_col   = tpl.get('body_color', '#374151')
    mt         = tpl.get('margin_top', 0.75)
    mb         = tpl.get('margin_bottom', 0.75)
    ml         = tpl.get('margin_left', 0.75)
    mr         = tpl.get('margin_right', 0.75)
    lh         = tpl.get('line_height', 1.3)
    para_sp    = tpl.get('para_spacing', 5)
    sec_sp     = tpl.get('section_spacing', 10)
    hdr_style  = tpl.get('header_style', 'underline')
    bullet     = tpl.get('bullet_char', '•')
    # The PDF engine (xhtml2pdf) can only render glyphs in its core WinAnsi fonts.
    # Coerce exotic bullets (◦ ▪ ‣ ○ …) to '•' so the PDF doesn't show a ■ box and the
    # preview (which uses this same renderer) stays identical to the download.
    if bullet not in ('•', '-', '*', '·', '–'):
        bullet = '•'

    def section_header(title: str) -> str:
        t = _esc(title.upper())
        base = (
            f'font-size:{hdr_sz}pt;font-weight:bold;font-family:{font};'
            f'color:{accent};letter-spacing:0.8px;display:block;'
        )
        margin = f'margin-top:{sec_sp}pt;margin-bottom:4pt;'

        if hdr_style == 'bar':
            return (
                f'<div style="{margin}background:{accent};padding:3px 6px;">'
                f'<span style="font-size:{hdr_sz}pt;font-weight:bold;'
                f'font-family:{font};color:#fff;letter-spacing:0.8px;">{t}</span></div>'
            )
        elif hdr_style == 'caps':
            return (
                f'<div style="{margin}border-bottom:1px solid {accent};padding-bottom:2px;">'
                f'<span style="{base}letter-spacing:1.8px;">{t}</span></div>'
            )
        elif hdr_style == 'plain':
            return (
                f'<div style="{margin}">'
                f'<span style="font-size:{hdr_sz}pt;font-weight:bold;font-family:{font};'
                f'color:{body_col};">{t}</span></div>'
            )
        else:  # underline (default)
            return (
                f'<div style="{margin}border-bottom:1.5px solid {accent};padding-bottom:2px;">'
                f'<span style="{base}">{t}</span></div>'
            )

    body_style = f'font-family:{font};font-size:{body_sz}pt;color:{body_col};line-height:{lh};'
    skills_layout = tpl.get('skills_layout', 'categorized')

    # ── Per-section builders (return '' when there's nothing to show) ──
    def build_summary(label):
        if not sections.get('summary'):
            return ''
        return (section_header(label) +
                f'<p style="{body_style}margin:{para_sp}pt 0 0 0;">{_rich(sections["summary"])}</p>')

    def build_skills(label):
        sk_list = [s for s in (sections.get('skills') or []) if s.get('items') or s.get('category')]
        if not sk_list:
            return ''
        html = section_header(label)
        if skills_layout == 'inline':
            joined = ', '.join(s.get('items', '').strip() for s in sk_list if s.get('items'))
            html += f'<div style="{body_style}margin-bottom:{para_sp // 2}pt;">{_rich(joined)}</div>'
        else:
            for sk in sk_list:
                cat   = _esc(sk.get('category', ''))
                items = _rich(sk.get('items', ''))
                lbl   = f'<strong>{cat}:</strong> ' if cat else ''
                html += (f'<div style="{body_style}margin-bottom:{para_sp // 2}pt;">'
                         f'{lbl}{items}</div>')
        return html

    def build_experience(label):
        exp_list = [e for e in (sections.get('experience') or []) if e.get('title')]
        if not exp_list:
            return ''
        html = section_header(label)
        for exp in exp_list:
            company_line = _esc(exp.get('company', ''))
            if exp.get('dates'):
                company_line += f' &nbsp;|&nbsp; {_esc(exp["dates"])}'
            html += (
                f'<div style="margin-bottom:{para_sp}pt;">'
                f'<div style="font-size:{body_sz + 1}pt;font-weight:bold;font-family:{font};'
                f'color:{name_col};">{_esc(exp.get("title", ""))}</div>'
                f'<div style="{body_style}margin-bottom:3pt;">{company_line}</div>'
            )
            for bl in [b for b in (exp.get('bullets') or []) if b]:
                html += (f'<div style="{body_style}padding-left:16pt;text-indent:-11pt;'
                         f'margin-bottom:2pt;">{_esc(bullet)}&nbsp;{_rich(bl)}</div>')
            html += '</div>'
        return html

    def build_education(label):
        edu_list = [e for e in (sections.get('education') or []) if e.get('degree')]
        if not edu_list:
            return ''
        html = section_header(label)
        for edu in edu_list:
            school = _esc(edu.get('school', ''))
            dates  = _esc(edu.get('dates', ''))
            line2  = f'{school} &nbsp;|&nbsp; {dates}' if (school and dates) else school or dates
            html += (f'<div style="{body_style}margin-bottom:{para_sp}pt;">'
                     f'<strong>{_esc(edu.get("degree", ""))}</strong>'
                     f'{"<br>" + line2 if line2 else ""}</div>')
        return html

    def build_certifications(label):
        cert_list = [c for c in (sections.get('certifications') or []) if c]
        if not cert_list:
            return ''
        html = section_header(label)
        for cert in cert_list:
            html += (f'<div style="{body_style}padding-left:16pt;text-indent:-11pt;'
                     f'margin-bottom:2pt;">{_esc(bullet)}&nbsp;{_rich(cert)}</div>')
        return html

    builders = {
        'summary': build_summary, 'skills': build_skills, 'experience': build_experience,
        'education': build_education, 'certifications': build_certifications,
    }

    parts: list[str] = []

    # ── Personal header (always first) ───────────────────────────────
    parts.append(
        f'<div style="text-align:center;margin-bottom:6pt;">'
        f'<div style="font-size:{name_sz}pt;font-weight:bold;font-family:{font};'
        f'color:{name_col};letter-spacing:0.5px;">{_esc(sections.get("name", ""))}</div>'
        f'<div style="font-size:{contact_sz}pt;font-family:{font};color:{body_col};margin-top:3pt;">'
        f'{_esc(sections.get("contact", ""))}</div>'
        f'</div>'
    )

    # ── Body sections in the configured order / visibility / labels ──
    layout = tpl.get('sections_layout') or DEFAULT_SECTION_LAYOUT
    for item in layout:
        if not item.get('visible', True):
            continue
        build = builders.get(item.get('key'))
        if not build:
            continue
        block = build(item.get('label') or item.get('key', '').title())
        if block:
            parts.append(block)

    inner = '\n'.join(parts)

    if not for_print:
        return inner

    page_css = (
        f'@page {{ size: letter; margin: {mt}in {mr}in {mb}in {ml}in; }}'
        f'body {{ margin:0;padding:0;background:#fff; }}'
    )
    return (
        f'<!DOCTYPE html><html><head>'
        f'<meta charset="utf-8">'
        f'<style>{page_css}</style>'
        f'</head><body>'
        f'<div style="font-family:{font};font-size:{body_sz}pt;color:{body_col};'
        f'line-height:{lh};background:#fff;">'
        f'{inner}'
        f'</div></body></html>'
    )


# ─── PDF export ──────────────────────────────────────────────────────────────

def export_pdf_html(sections: dict, tpl: dict) -> str:
    """
    Return a fully self-contained HTML string optimised for browser print-to-PDF.
    Includes an auto-print script so opening it in the browser immediately
    triggers the print dialog.
    """
    inner = render_resume_html(sections, tpl, for_print=False)

    mt = tpl.get('margin_top',    0.75)
    mb = tpl.get('margin_bottom', 0.75)
    ml = tpl.get('margin_left',   0.75)
    mr = tpl.get('margin_right',  0.75)
    font = tpl.get('font_family', 'Georgia, serif')
    body_sz = tpl.get('body_size', 11)
    body_col = tpl.get('body_color', '#374151')
    lh = tpl.get('line_height', 1.3)

    page_css = (
        f'@page {{ size: letter; margin: {mt}in {mr}in {mb}in {ml}in; }}'
        f'* {{ box-sizing: border-box; }}'
        f'body {{ margin:0; padding:0; background:#fff; '
        f'font-family:{font}; font-size:{body_sz}pt; '
        f'color:{body_col}; line-height:{lh}; }}'
        f'@media screen {{ body {{ padding: 0.5in; max-width: 8.5in; margin: 0 auto; }} }}'
    )

    return (
        f'<!DOCTYPE html><html lang="en"><head>'
        f'<meta charset="utf-8">'
        f'<title>Resume</title>'
        f'<style>{page_css}</style>'
        f'</head><body>'
        f'{inner}'
        f'<script>window.onload = function() {{ window.print(); }}</script>'
        f'</body></html>'
    )


def export_pdf(sections: dict, tpl: dict) -> bytes:
    """
    Try WeasyPrint → xhtml2pdf → raise ImportError with helpful message.
    Callers should catch ImportError and fall back to export_pdf_html().
    """
    # Try WeasyPrint
    try:
        import weasyprint  # type: ignore
        html = render_resume_html(sections, tpl, for_print=True)
        pdf  = weasyprint.HTML(string=html).write_pdf()
        return pdf
    except ImportError:
        pass
    except Exception as e:
        raise ValueError(f"WeasyPrint error: {e}") from e

    # Try xhtml2pdf
    try:
        from xhtml2pdf import pisa  # type: ignore
        html   = render_resume_html(sections, tpl, for_print=True)
        buf    = BytesIO()
        status = pisa.CreatePDF(html, dest=buf, encoding='utf-8')
        if status.err:
            raise ValueError(f"xhtml2pdf error: {status.err}")
        return buf.getvalue()
    except ImportError:
        pass

    raise ImportError("pdf_libs_missing")


# ─── DOCX export ─────────────────────────────────────────────────────────────

def export_docx(sections: dict, tpl: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    font_name   = tpl.get('font_family', 'Georgia').split(',')[0].strip().strip('"\'')
    name_sz     = tpl.get('name_size', 22)
    hdr_sz      = tpl.get('header_size', 13)
    body_sz     = tpl.get('body_size', 11)
    contact_sz  = tpl.get('contact_size', 10)
    accent_rgb  = RGBColor(*_hex_to_rgb(tpl.get('accent_color', '#1e3a5f')))
    name_rgb    = RGBColor(*_hex_to_rgb(tpl.get('name_color', '#111827')))
    body_rgb    = RGBColor(*_hex_to_rgb(tpl.get('body_color', '#374151')))
    mt          = tpl.get('margin_top', 0.75)
    mb          = tpl.get('margin_bottom', 0.75)
    ml          = tpl.get('margin_left', 0.75)
    mr          = tpl.get('margin_right', 0.75)
    para_sp     = Pt(tpl.get('para_spacing', 5))
    bullet_chr  = tpl.get('bullet_char', '•')

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Inches(mt)
    sec.bottom_margin = Inches(mb)
    sec.left_margin   = Inches(ml)
    sec.right_margin  = Inches(mr)

    # Remove default paragraph spacing from Normal style
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)

    def _run(para, text, bold=False, size=None, color=None, italic=False):
        run = para.add_run(str(text))
        run.bold   = bold
        run.italic = italic
        run.font.name = font_name
        run.font.size = Pt(size or body_sz)
        if color:
            run.font.color.rgb = color
        return run

    def _run_rich(para, text, size=None, color=None, prefix=''):
        """Add runs honoring **bold** markers (so DOCX matches the bolded preview/PDF)."""
        if prefix:
            _run(para, prefix, size=size, color=color)
        for seg, is_bold in _bold_segments(text):
            if seg:
                _run(para, seg, bold=is_bold, size=size, color=color)

    def _add_section_header(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        _run(p, title.upper(), bold=True, size=hdr_sz, color=accent_rgb)
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), tpl.get('accent_color', '#1e3a5f').lstrip('#'))
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    # ── Name ─────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _run(p, sections.get('name', ''), bold=True, size=name_sz, color=name_rgb)

    # ── Contact ──────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    _run(p, sections.get('contact', ''), size=contact_sz)

    skills_layout = tpl.get('skills_layout', 'categorized')

    # ── Per-section builders (honor the configured order/labels/visibility) ──
    def docx_summary(label):
        if not sections.get('summary'):
            return
        _add_section_header(label)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = para_sp
        _run_rich(p, sections['summary'])

    def docx_skills(label):
        sk_list = [s for s in (sections.get('skills') or []) if s.get('items') or s.get('category')]
        if not sk_list:
            return
        _add_section_header(label)
        if skills_layout == 'inline':
            joined = ', '.join(s.get('items', '').strip() for s in sk_list if s.get('items'))
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _run_rich(p, joined)
        else:
            for sk in sk_list:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                if sk.get('category'):
                    _run(p, sk['category'] + ': ', bold=True)
                _run_rich(p, sk.get('items', ''))

    def docx_experience(label):
        exp_list = [e for e in (sections.get('experience') or []) if e.get('title')]
        if not exp_list:
            return
        _add_section_header(label)
        for exp in exp_list:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(1)
            _run(p, exp.get('title', ''), bold=True, size=body_sz + 1, color=name_rgb)

            company = exp.get('company', '')
            dates   = exp.get('dates', '')
            line2   = f'{company}  |  {dates}' if (company and dates) else company or dates
            if line2:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(2)
                _run(p2, line2)

            for bl in [b for b in (exp.get('bullets') or []) if b]:
                p3 = doc.add_paragraph()
                p3.paragraph_format.left_indent = Inches(0.15)
                p3.paragraph_format.space_after = Pt(1)
                _run_rich(p3, bl, prefix=f'{bullet_chr}  ')

    def docx_education(label):
        edu_list = [e for e in (sections.get('education') or []) if e.get('degree')]
        if not edu_list:
            return
        _add_section_header(label)
        for edu in edu_list:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _run(p, edu.get('degree', ''), bold=True)
            school = edu.get('school', '')
            dates  = edu.get('dates', '')
            line2  = f'{school}  |  {dates}' if (school and dates) else school or dates
            if line2:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = para_sp
                _run(p2, line2)

    def docx_certs(label):
        cert_list = [c for c in (sections.get('certifications') or []) if c]
        if not cert_list:
            return
        _add_section_header(label)
        for cert in cert_list:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(1)
            _run_rich(p, cert, prefix=f'{bullet_chr}  ')

    docx_builders = {
        'summary': docx_summary, 'skills': docx_skills, 'experience': docx_experience,
        'education': docx_education, 'certifications': docx_certs,
    }
    for item in (tpl.get('sections_layout') or DEFAULT_SECTION_LAYOUT):
        if not item.get('visible', True):
            continue
        b = docx_builders.get(item.get('key'))
        if b:
            b(item.get('label') or item.get('key', '').title())

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
